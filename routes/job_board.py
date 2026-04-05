import io
import logging
from datetime import datetime
from functools import wraps

from flask import Blueprint, jsonify, request, send_file, session

from extensions import db
from models.job_post import JobPost
from models.settings import Setting

logger = logging.getLogger(__name__)
job_board_bp = Blueprint('job_board', __name__, url_prefix='/api/jobboard')


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('admin_logged_in'):
            return jsonify({'error': 'Unauthorized'}), 401
        return f(*args, **kwargs)
    return decorated


def _get_groq_client():
    import groq
    key = Setting.get('groq_api_key', '')
    if not key:
        raise ValueError('Groq API key not configured.')
    return groq.Groq(api_key=key)


# ── Public API ──────────────────────────────────────────────────────────────

@job_board_bp.get('/published')
def public_list():
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 12))
    search = request.args.get('q', '').strip().lower()
    job_type = request.args.get('type', '').strip().lower()
    tag = request.args.get('tag', '').strip().lower()

    q = JobPost.query.filter_by(status='published').order_by(JobPost.featured.desc(), JobPost.updated_at.desc())
    posts = q.all()

    if search:
        posts = [p for p in posts if search in (p.title or '').lower()
                 or search in (p.company or '').lower()
                 or search in (p.location or '').lower()
                 or search in (p.tags or '').lower()
                 or search in (p.description or '').lower()
                 or search in (p.original_description or '').lower()]
    if job_type:
        posts = [p for p in posts if job_type in (p.job_type or '').lower()]
    if tag:
        posts = [p for p in posts if tag in (p.tags or '').lower()]

    total = len(posts)
    start = (page - 1) * per_page
    page_posts = posts[start:start + per_page]
    return jsonify({
        'total': total,
        'page': page,
        'pages': max(1, -(-total // per_page)),
        'posts': [p.to_dict() for p in page_posts],
    })


@job_board_bp.get('/published/<int:post_id>')
def public_detail(post_id):
    post = JobPost.query.filter_by(id=post_id, status='published').first_or_404()
    return jsonify(post.to_dict())


def _ai_rewrite_job(post_dict: dict) -> str:
    """Rewrite a single job description via AI. Returns rewritten text or original on failure."""
    from utils.ai_engine import rewrite_job_description
    src = (post_dict.get('original_description') or post_dict.get('description') or '').strip()
    if len(src) < 80:
        return src
    try:
        return rewrite_job_description(
            title=post_dict.get('title', ''),
            company=post_dict.get('company', ''),
            raw_description=src,
        )
    except Exception as e:
        logger.warning('AI rewrite skipped (%s – %s): %s', post_dict.get('title', '?'), post_dict.get('source', '?'), e)
        return src


@job_board_bp.post('/auto-rewrite')
@admin_required
def auto_rewrite():
    """
    Background batch-rewrite endpoint.
    Finds up to `batch` published jobs that haven't been AI-rewritten yet,
    rewrites them concurrently, saves to DB, and returns progress info.
    """
    import concurrent.futures
    batch = min(int(request.json.get('batch', 5)) if request.is_json else 5, 10)

    pending = JobPost.query.filter_by(status='published', ai_rewritten=False).limit(batch).all()
    total_remaining = JobPost.query.filter_by(status='published', ai_rewritten=False).count()

    if not pending:
        return jsonify({'done': 0, 'remaining': 0, 'finished': True})

    def _rewrite_one(post):
        src = (post.original_description or post.description or '').strip()
        if len(src) < 80:
            post.ai_rewritten = True
            return post, None
        try:
            from utils.ai_engine import rewrite_job_description
            rewritten = rewrite_job_description(
                title=post.title or '',
                company=post.company or '',
                raw_description=src,
            )
            return post, rewritten
        except Exception as e:
            logger.warning('Auto-rewrite failed for job %s: %s', post.id, e)
            return post, None

    with concurrent.futures.ThreadPoolExecutor(max_workers=min(batch, 5)) as ex:
        futures = {ex.submit(_rewrite_one, p): p for p in pending}
        results = []
        for f in concurrent.futures.as_completed(futures, timeout=60):
            try:
                results.append(f.result())
            except Exception:
                pass

    done = 0
    for post, rewritten in results:
        if rewritten:
            post.description = rewritten
        post.ai_rewritten = True
        post.updated_at = datetime.utcnow()
        done += 1

    db.session.commit()
    still_remaining = JobPost.query.filter_by(status='published', ai_rewritten=False).count()
    return jsonify({'done': done, 'remaining': still_remaining, 'finished': still_remaining == 0})


@job_board_bp.get('/rewrite-status')
@admin_required
def rewrite_status():
    """Return count of jobs still pending AI rewrite."""
    pending = JobPost.query.filter_by(status='published', ai_rewritten=False).count()
    total = JobPost.query.filter_by(status='published').count()
    return jsonify({'pending': pending, 'total': total, 'finished': pending == 0})


@job_board_bp.get('/live-search')
def live_search():
    """Search local DB + live external APIs simultaneously, return merged results."""
    import concurrent.futures
    from utils.job_aggregator import fetch_remotive, fetch_arbeitnow, fetch_remoteok

    q = request.args.get('q', '').strip()
    job_type = request.args.get('type', '').strip().lower()
    search = q.lower()

    # --- Local DB results (instant) ---
    db_posts = JobPost.query.filter_by(status='published').order_by(
        JobPost.featured.desc(), JobPost.updated_at.desc()).all()
    if search:
        db_posts = [p for p in db_posts if search in (p.title or '').lower()
                    or search in (p.company or '').lower()
                    or search in (p.location or '').lower()
                    or search in (p.tags or '').lower()
                    or search in (p.description or '').lower()
                    or search in (p.original_description or '').lower()]
    if job_type:
        db_posts = [p for p in db_posts if job_type in (p.job_type or '').lower()]

    local_results = []
    for p in db_posts:
        d = p.to_dict()
        d['is_live'] = False
        local_results.append(d)

    # --- Live external API results + AI rewrite (parallel) ---
    live_results = []
    if q:
        def _remotive():
            try:
                return fetch_remotive(search=q, limit=10)
            except Exception:
                return []

        def _arbeitnow():
            try:
                raw = fetch_arbeitnow(limit=30)
                return [r for r in raw if search in (r.get('title') or '').lower()
                        or search in (r.get('company') or '').lower()
                        or search in (r.get('tags') or '').lower()
                        or search in (r.get('original_description') or '').lower()][:10]
            except Exception:
                return []

        def _remoteok():
            try:
                raw = fetch_remoteok(limit=50)
                return [r for r in raw if search in (r.get('title') or '').lower()
                        or search in (r.get('company') or '').lower()
                        or search in (r.get('tags') or '').lower()
                        or search in (r.get('original_description') or '').lower()][:10]
            except Exception:
                return []

        raw_live = []
        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as ex:
            fetch_futures = [ex.submit(_remotive), ex.submit(_arbeitnow), ex.submit(_remoteok)]
            for f in concurrent.futures.as_completed(fetch_futures, timeout=12):
                try:
                    raw_live.extend(f.result())
                except Exception:
                    pass

        # Deduplicate live vs local by external_id / title+company
        local_ext_ids = {p.external_id for p in db_posts if p.external_id}
        local_keys = {(p.title or '').lower() + '|' + (p.company or '').lower() for p in db_posts}
        seen_live = set()
        deduped_live = []
        for r in raw_live:
            ext_id = r.get('external_id', '')
            key = (r.get('title') or '').lower() + '|' + (r.get('company') or '').lower()
            if ext_id and ext_id in local_ext_ids:
                continue
            if key in local_keys or key in seen_live:
                continue
            seen_live.add(key)
            r['is_live'] = True
            r['id'] = None
            deduped_live.append(r)

        if job_type:
            deduped_live = [r for r in deduped_live if job_type in (r.get('job_type') or '').lower()]

        # AI-rewrite live job descriptions concurrently
        if deduped_live:
            with concurrent.futures.ThreadPoolExecutor(max_workers=min(len(deduped_live), 8)) as ex:
                rewrite_futures = {ex.submit(_ai_rewrite_job, job): i for i, job in enumerate(deduped_live)}
                for f, idx in rewrite_futures.items():
                    try:
                        deduped_live[idx]['description'] = f.result(timeout=15)
                        deduped_live[idx]['ai_rewritten'] = True
                    except Exception:
                        pass

        live_results = deduped_live

    all_results = local_results + live_results
    return jsonify({
        'total': len(all_results),
        'local_count': len(local_results),
        'live_count': len(live_results),
        'results': all_results,
    })


@job_board_bp.get('/featured')
def featured_posts():
    posts = JobPost.query.filter_by(status='published', featured=True).order_by(JobPost.updated_at.desc()).limit(6).all()
    if not posts:
        posts = JobPost.query.filter_by(status='published').order_by(JobPost.updated_at.desc()).limit(6).all()
    return jsonify([p.to_dict() for p in posts])


# ── Export API ──────────────────────────────────────────────────────────────

@job_board_bp.post('/export')
@admin_required
def export_jobs():
    data = request.json or {}
    ids = data.get('ids', [])
    fmt = data.get('format', 'txt').lower()
    search = data.get('search', '').strip().lower()
    job_type = data.get('job_type', '').strip().lower()

    if ids:
        posts = JobPost.query.filter(JobPost.id.in_(ids), JobPost.status == 'published').all()
    else:
        q = JobPost.query.filter_by(status='published').order_by(JobPost.featured.desc(), JobPost.updated_at.desc())
        posts = q.all()
        if search:
            posts = [p for p in posts if search in (p.title or '').lower()
                     or search in (p.company or '').lower()
                     or search in (p.location or '').lower()
                     or search in (p.tags or '').lower()
                     or search in (p.description or '').lower()
                     or search in (p.original_description or '').lower()]
        if job_type:
            posts = [p for p in posts if job_type in (p.job_type or '').lower()]

    if not posts:
        return jsonify({'error': 'No jobs to export'}), 400

    if fmt == 'txt':
        lines = []
        for p in posts:
            lines.append('=' * 60)
            lines.append(f'Title:    {p.title}')
            lines.append(f'Company:  {p.company or "N/A"}')
            lines.append(f'Location: {p.location or "N/A"}')
            lines.append(f'Type:     {p.job_type or "N/A"}')
            if p.salary:
                lines.append(f'Salary:   {p.salary}')
            if p.apply_url:
                lines.append(f'Apply:    {p.apply_url}')
            if p.tags:
                lines.append(f'Tags:     {p.tags}')
            desc = (p.description or p.original_description or '').strip()
            if desc:
                lines.append('')
                lines.append(desc[:800])
            lines.append('')
        content = '\n'.join(lines).encode('utf-8')
        buf = io.BytesIO(content)
        buf.seek(0)
        return send_file(buf, mimetype='text/plain',
                         as_attachment=True, download_name='job_listings.txt')

    elif fmt == 'docx':
        from docx import Document
        from docx.shared import Pt, RGBColor
        doc = Document()
        doc.add_heading('Job Listings Export', 0)
        for p in posts:
            h = doc.add_heading(p.title, level=1)
            h.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            details = []
            if p.company:  details.append(f'Company: {p.company}')
            if p.location: details.append(f'Location: {p.location}')
            if p.job_type: details.append(f'Type: {p.job_type}')
            if p.salary:   details.append(f'Salary: {p.salary}')
            if p.apply_url: details.append(f'Apply: {p.apply_url}')
            if p.tags:      details.append(f'Tags: {p.tags}')
            for d in details:
                para = doc.add_paragraph(d)
                para.runs[0].bold = True
            desc = (p.description or p.original_description or '').strip()
            if desc:
                doc.add_paragraph(desc[:800])
            doc.add_paragraph('')
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name='job_listings.docx')

    elif fmt == 'pdf':
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=20*mm, rightMargin=20*mm,
                                topMargin=20*mm, bottomMargin=20*mm)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('jb_title', parent=styles['Heading1'],
                                     textColor=colors.HexColor('#4F46E5'), fontSize=14, spaceAfter=4)
        meta_style = ParagraphStyle('jb_meta', parent=styles['Normal'],
                                    textColor=colors.HexColor('#475569'), fontSize=9, spaceAfter=2)
        desc_style = ParagraphStyle('jb_desc', parent=styles['Normal'],
                                    fontSize=9, leading=13, spaceAfter=6,
                                    textColor=colors.HexColor('#334155'))
        story = [Paragraph('Job Listings Export', styles['Title']), Spacer(1, 8*mm)]
        for p in posts:
            story.append(Paragraph(p.title or 'Untitled', title_style))
            meta_parts = []
            if p.company:   meta_parts.append(f'<b>Company:</b> {p.company}')
            if p.location:  meta_parts.append(f'<b>Location:</b> {p.location}')
            if p.job_type:  meta_parts.append(f'<b>Type:</b> {p.job_type}')
            if p.salary:    meta_parts.append(f'<b>Salary:</b> {p.salary}')
            if p.apply_url: meta_parts.append(f'<b>Apply:</b> {p.apply_url}')
            if p.tags:      meta_parts.append(f'<b>Tags:</b> {p.tags}')
            for m in meta_parts:
                story.append(Paragraph(m, meta_style))
            desc = (p.description or p.original_description or '').strip()
            if desc:
                safe_desc = desc[:800].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_desc, desc_style))
            story.append(HRFlowable(width='100%', thickness=0.5, color=colors.HexColor('#E2E8F0')))
            story.append(Spacer(1, 5*mm))
        doc.build(story)
        buf.seek(0)
        return send_file(buf, mimetype='application/pdf',
                         as_attachment=True, download_name='job_listings.pdf')

    return jsonify({'error': 'Invalid format'}), 400


# ── Admin API ───────────────────────────────────────────────────────────────

@job_board_bp.get('/admin/posts')
@admin_required
def admin_list():
    status = request.args.get('status', 'all')
    q = JobPost.query.order_by(JobPost.created_at.desc())
    if status != 'all':
        q = q.filter_by(status=status)
    posts = q.all()

    counts = {
        'draft': JobPost.query.filter_by(status='draft').count(),
        'published': JobPost.query.filter_by(status='published').count(),
        'archived': JobPost.query.filter_by(status='archived').count(),
        'total': JobPost.query.count(),
    }
    return jsonify({'posts': [p.to_dict() for p in posts], 'counts': counts})


@job_board_bp.get('/admin/posts/<int:post_id>')
@admin_required
def admin_get(post_id):
    post = JobPost.query.get_or_404(post_id)
    d = post.to_dict()
    d['original_description'] = post.original_description
    d['description_raw'] = post.description
    return jsonify(d)


@job_board_bp.put('/admin/posts/<int:post_id>')
@admin_required
def admin_update(post_id):
    post = JobPost.query.get_or_404(post_id)
    data = request.json or {}
    for field in ['title', 'company', 'location', 'job_type', 'salary', 'apply_url', 'description']:
        if field in data:
            setattr(post, field, data[field])
    if 'tags' in data:
        tags = data['tags']
        post.tags = ', '.join(tags) if isinstance(tags, list) else tags
    if 'featured' in data:
        post.featured = bool(data['featured'])
    if 'status' in data and data['status'] in ('draft', 'published', 'archived'):
        post.status = data['status']
    post.updated_at = datetime.utcnow()
    db.session.commit()
    return jsonify({'success': True, 'post': post.to_dict()})


@job_board_bp.delete('/admin/posts/<int:post_id>')
@admin_required
def admin_delete(post_id):
    post = JobPost.query.get_or_404(post_id)
    db.session.delete(post)
    db.session.commit()
    return jsonify({'success': True})


@job_board_bp.post('/admin/posts/bulk')
@admin_required
def admin_bulk():
    data = request.json or {}
    action = data.get('action')
    ids = data.get('ids', [])
    if not ids or action not in ('publish', 'archive', 'draft', 'delete'):
        return jsonify({'success': False, 'error': 'Invalid request'}), 400
    posts = JobPost.query.filter(JobPost.id.in_(ids)).all()
    if action == 'delete':
        for p in posts:
            db.session.delete(p)
    else:
        for p in posts:
            p.status = action
            p.updated_at = datetime.utcnow()
    db.session.commit()
    return jsonify({'success': True, 'affected': len(posts)})


@job_board_bp.post('/admin/posts/<int:post_id>/rewrite')
@admin_required
def admin_rewrite(post_id):
    post = JobPost.query.get_or_404(post_id)
    source_text = post.original_description or post.description
    if not source_text or len(source_text.strip()) < 50:
        return jsonify({'success': False, 'error': 'Job description is too short to rewrite.'}), 400
    try:
        from utils.ai_engine import rewrite_job_description
        rewritten = rewrite_job_description(
            title=post.title or '',
            company=post.company or '',
            raw_description=source_text,
        )
        post.description = rewritten
        post.ai_rewritten = True
        post.updated_at = datetime.utcnow()
        db.session.commit()
        return jsonify({'success': True, 'description': rewritten})
    except Exception as e:
        logger.error('Groq rewrite error: %s', e)
        return jsonify({'success': False, 'error': str(e)}), 500


@job_board_bp.post('/admin/fetch')
@admin_required
def admin_fetch():
    data = request.json or {}
    sources = data.get('sources', ['remotive', 'arbeitnow', 'remoteok'])
    search = data.get('search', '')
    limit = int(data.get('limit_per_source', 15))

    adzuna_id = Setting.get('adzuna_app_id', '')
    adzuna_key = Setting.get('adzuna_app_key', '')

    try:
        from utils.job_aggregator import aggregate
        posts_data = aggregate(
            sources=sources,
            search=search,
            limit_per_source=limit,
            adzuna_app_id=adzuna_id or None,
            adzuna_app_key=adzuna_key or None,
        )
    except Exception as e:
        logger.error('Aggregation error: %s', e)
        return jsonify({'success': False, 'error': str(e)}), 500

    added = 0
    skipped = 0
    for p in posts_data:
        ext_id = p.get('external_id', '')
        if ext_id and JobPost.query.filter_by(external_id=ext_id).first():
            skipped += 1
            continue
        desc = p.get('original_description', '')
        post = JobPost(
            external_id=ext_id,
            source=p.get('source', 'unknown'),
            title=p.get('title', ''),
            company=p.get('company', ''),
            location=p.get('location', ''),
            job_type=p.get('job_type', ''),
            salary=p.get('salary', ''),
            tags=p.get('tags', ''),
            apply_url=p.get('apply_url', ''),
            original_description=desc,
            description=desc,
            status='draft',
        )
        db.session.add(post)
        added += 1

    db.session.commit()
    pending_rewrite = JobPost.query.filter_by(status='published', ai_rewritten=False).count()
    return jsonify({
        'success': True,
        'added': added,
        'skipped': skipped,
        'total': len(posts_data),
        'pending_rewrite': pending_rewrite,
    })
